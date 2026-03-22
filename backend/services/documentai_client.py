"""
Denial letter parsing pipeline.

  1. Detect MIME type from filename/content
  2. Convert non-native formats (DOCX → text via python-docx, TXT → direct)
  3. Send PDFs/images to Google Document AI OCR for text extraction
  4. Pass extracted text to Gemini Flash for structured field parsing
"""
import io
import json
import mimetypes
import os
import re
import subprocess
import tempfile
from pathlib import Path

from google.api_core.client_options import ClientOptions
from google.cloud import documentai

import google.generativeai as genai

from config import settings

# ─── Clients ──────────────────────────────────────────────────────────────────

# Resolve credentials path relative to this file's directory so it works
# regardless of where the server is launched from.
_creds_path = settings.google_application_credentials
if _creds_path:
    _abs_creds = str(Path(__file__).parent.parent / _creds_path.lstrip("./"))
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = _abs_creds

genai.configure(api_key=settings.google_api_key)
_gemini = genai.GenerativeModel("gemini-2.5-flash")

_PROCESSOR_NAME = (
    f"projects/{settings.google_project_number}"
    f"/locations/{settings.documentai_location}"
    f"/processors/{settings.documentai_processor_id}"
)

# Singleton Document AI client — initialised once at startup, not per request.
_documentai_client: documentai.DocumentProcessorServiceClient | None = None

def _get_documentai_client() -> documentai.DocumentProcessorServiceClient:
    global _documentai_client
    if _documentai_client is None:
        _documentai_client = documentai.DocumentProcessorServiceClient(
            client_options=ClientOptions(
                api_endpoint=f"{settings.documentai_location}-documentai.googleapis.com"
            )
        )
    return _documentai_client

# Document AI supports these MIME types natively
_NATIVE_MIME_TYPES = {
    "application/pdf",
    "image/tiff",
    "image/gif",
    "image/jpeg",
    "image/png",
    "image/bmp",
    "image/webp",
}

# ─── MIME detection ───────────────────────────────────────────────────────────

def _detect_mime(filename: str, content: bytes) -> str:
    mime, _ = mimetypes.guess_type(filename)
    if mime:
        return mime
    if content[:4] == b"%PDF":
        return "application/pdf"
    if content[:2] == b"PK":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "application/msword"
    return "application/octet-stream"


# ─── Format converters ────────────────────────────────────────────────────────

def _extract_docx_text(content: bytes) -> str:
    """Extract plain text from DOCX using python-docx."""
    import docx
    doc = docx.Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes → PDF bytes using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "input.docx"
        pdf_path = Path(tmpdir) / "input.pdf"
        docx_path.write_bytes(docx_bytes)
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, str(docx_path)],
                check=True, capture_output=True, timeout=30,
            )
            return pdf_path.read_bytes()
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            raise RuntimeError(
                "LibreOffice is required to convert DOCX files. "
                "Install with: apt-get install libreoffice  |  brew install libreoffice"
            ) from e


# ─── Document AI OCR ─────────────────────────────────────────────────────────

def _ocr_with_documentai(file_bytes: bytes, mime_type: str) -> str:
    """Send document bytes to Document AI OCR and return full extracted text."""
    client = _get_documentai_client()
    raw_document = documentai.RawDocument(content=file_bytes, mime_type=mime_type)
    request = documentai.ProcessRequest(name=_PROCESSOR_NAME, raw_document=raw_document)
    result = client.process_document(request=request)
    return result.document.text


# ─── Gemini structured extraction ─────────────────────────────────────────────

_EXTRACTION_PROMPT = """\
You are a medical insurance claims specialist. Extract structured information from
this insurance denial letter text. The letter may be from any commercial insurer,
Medicare Advantage plan, or Medicaid managed care organization.

Extract these fields:
- patient_name: Full name of the patient (or "Unknown" if not found)
- claim_id: Claim number, reference number, or authorization number (or "Unknown")
- denial_reason: The primary stated reason for denial — be specific, quote the letter language
- cpb_code_cited: Any CPB number, LCD number, NCD number, policy number, or coverage
  determination number cited as justification (or "Not cited")
- deadline: The deadline to file an appeal in ISO date format YYYY-MM-DD
  (or "Not specified" if absent)
- required_docs: Array of specific documents or information the insurer says would
  support an appeal or reversal — extract every item mentioned
- insurer_name: Name of the insurance company or plan
- denial_date: Date of the denial letter in YYYY-MM-DD format (or "Unknown")
- service_denied: The procedure, drug, device, or service that was denied
- denial_codes: Array of any denial/remark codes cited (e.g. CO-50, PR-96, N30)

Respond in valid JSON only. No markdown. No preamble.

DENIAL LETTER TEXT:
{text}
"""


_TLDR_PROMPT = """\
You are explaining a health insurance denial to a stressed patient in plain English.
Write a TL;DR of 2-3 sentences max. Be direct, empathetic, and actionable.
No jargon. No bullet points. Just clear prose.

Format: "[Insurer] denied your [service] because [short reason]. \
You have until [deadline] to appeal — [X days away if known]. \
[One sentence on why this is beatable and what the key next step is.]"

Use the structured data below. If a field is unknown, skip it gracefully.

{fields}
"""


async def extract_denial_fields(raw_text: str) -> dict:
    """Use Gemini Flash to parse structured fields from denial letter text."""
    prompt = _EXTRACTION_PROMPT.format(text=raw_text)
    response = _gemini.generate_content(prompt)
    result = response.text.strip()
    result = re.sub(r"^```(?:json)?\s*", "", result)
    result = re.sub(r"\s*```$", "", result)
    return json.loads(result)


async def generate_tldr(fields: dict) -> str:
    """Generate a plain-English TL;DR from extracted denial fields."""
    summary = {k: v for k, v in fields.items() if k != "raw_text"}
    response = _gemini.generate_content(_TLDR_PROMPT.format(fields=json.dumps(summary, indent=2)))
    return response.text.strip()


# ─── Main entry point ─────────────────────────────────────────────────────────

async def parse_denial_document(filename: str, content: bytes) -> dict:
    """
    Full pipeline: detect format → extract text → structured extraction.
    """
    mime = _detect_mime(filename, content)
    ext = Path(filename).suffix.lower()

    raw_text: str | None = None

    # Plain text — skip OCR entirely
    if mime == "text/plain" or ext == ".txt":
        raw_text = content.decode("utf-8", errors="replace")

    # DOCX — extract text with python-docx (no LibreOffice needed)
    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == ".docx":
        try:
            raw_text = _extract_docx_text(content)
        except Exception:
            content = _docx_to_pdf_bytes(content)
            mime = "application/pdf"

    # Legacy DOC → PDF via LibreOffice
    elif mime == "application/msword" or ext == ".doc":
        content = _docx_to_pdf_bytes(content)
        mime = "application/pdf"

    # ODT, RTF, etc. — try LibreOffice, fall back to UTF-8 decode
    elif mime not in _NATIVE_MIME_TYPES:
        try:
            content = _docx_to_pdf_bytes(content)
            mime = "application/pdf"
        except RuntimeError:
            raw_text = content.decode("utf-8", errors="replace")

    # PDF / images → Document AI OCR
    if raw_text is None:
        raw_text = _ocr_with_documentai(content, mime)

    fields = await extract_denial_fields(raw_text)
    fields["raw_text"] = raw_text
    fields["tldr"] = await generate_tldr(fields)
    return fields
